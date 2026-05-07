"""
app/parsers/document_parser.py
───────────────────────────────
Universal document parser.

Supports:
  - PDF   (.pdf)     → via pdfplumber (text) + pytesseract fallback (scanned)
  - DOCX  (.docx)    → via python-docx
  - XLSX  (.xlsx)    → via openpyxl (tables → CSV-like text)
  - CSV   (.csv)     → plain text
  - TXT   (.txt)     → plain text
  - Images (.png, .jpg, .jpeg) → pytesseract OCR
  - HTML  (.html)    → html2text

The parsed text is always returned as a single UTF-8 string.
This module contains ONLY pure functions — no side effects, no I/O except reading the file.
"""
from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Union

logger = logging.getLogger("ai-agent.parsers")


def parse_document(file_bytes: bytes, filename: str) -> str:
    """
    Parse any supported document format into plain text.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename (used to determine MIME type by extension).

    Returns:
        Extracted text content as a single string.

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext == ".docx":
        return _parse_docx(file_bytes)
    elif ext in (".xlsx", ".xls"):
        return _parse_excel(file_bytes)
    elif ext == ".csv":
        return _parse_csv(file_bytes)
    elif ext in (".txt", ".md"):
        return file_bytes.decode("utf-8", errors="replace")
    elif ext in (".html", ".htm"):
        return _parse_html(file_bytes)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        return _parse_image_ocr(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: '{ext}'. Supported: pdf, docx, xlsx, csv, txt, html, png/jpg.")


# ─── Format-Specific Parsers ──────────────────────────────────────────────────


def _parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF using pdfplumber.
    Falls back to pytesseract OCR for scanned pages with no selectable text.
    """
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError("pdfplumber is required for PDF parsing. Install: pip install pdfplumber") from exc

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{page_text.strip()}")
            else:
                # Scanned page — attempt OCR
                logger.debug("Page %d has no selectable text; attempting OCR.", page_num)
                ocr_text = _ocr_pdf_page(page)
                if ocr_text.strip():
                    text_parts.append(f"--- Page {page_num} (OCR) ---\n{ocr_text.strip()}")

    result = "\n\n".join(text_parts)
    logger.info("PDF parsed: %d pages, %d chars.", len(text_parts), len(result))
    return result


def _ocr_pdf_page(page: object) -> str:  # page: pdfplumber.Page
    """Use pytesseract to OCR a single pdfplumber page."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("pytesseract/Pillow not installed — skipping OCR for scanned page.")
        return ""
    try:
        img = page.to_image(resolution=200).original  # type: ignore[attr-defined]
        return pytesseract.image_to_string(img)
    except Exception as exc:
        logger.warning("OCR failed for page: %s", exc)
        return ""


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx is required for DOCX parsing. Install: pip install python-docx") from exc

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    table_rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                table_rows.append(row_text)

    parts = paragraphs
    if table_rows:
        parts += ["\n=== TABLES ==="] + table_rows

    result = "\n".join(parts)
    logger.info("DOCX parsed: %d paragraphs, %d chars.", len(paragraphs), len(result))
    return result


def _parse_excel(file_bytes: bytes) -> str:
    """
    Extract text from Excel (.xlsx/.xls) by iterating all worksheets.
    Each row is rendered as a pipe-delimited line for the LLM to understand.
    """
    try:
        import openpyxl
    except ImportError as exc:
        raise ImportError("openpyxl is required for Excel parsing. Install: pip install openpyxl") from exc

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    all_parts: list[str] = []
    for sheet in wb.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_str.strip(" |"):
                rows.append(row_str)
        if rows:
            all_parts.append(f"=== Sheet: {sheet.title} ===\n" + "\n".join(rows))

    result = "\n\n".join(all_parts)
    logger.info("Excel parsed: %d sheets, %d chars.", len(wb.worksheets), len(result))
    return result


def _parse_csv(file_bytes: bytes) -> str:
    """Parse a CSV file into a readable text block."""
    import csv

    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [" | ".join(row) for row in reader if any(c.strip() for c in row)]
    result = "\n".join(rows)
    logger.info("CSV parsed: %d rows, %d chars.", len(rows), len(result))
    return result


def _parse_html(file_bytes: bytes) -> str:
    """Convert HTML to plain text using html2text."""
    try:
        import html2text
    except ImportError as exc:
        raise ImportError("html2text is required. Install: pip install html2text") from exc

    h = html2text.HTML2Text()
    h.ignore_links = True
    h.ignore_images = True
    return h.handle(file_bytes.decode("utf-8", errors="replace"))


def _parse_image_ocr(file_bytes: bytes) -> str:
    """OCR an image using pytesseract."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise ImportError("pytesseract and Pillow are required. Install: pip install pytesseract Pillow") from exc

    img = Image.open(io.BytesIO(file_bytes))
    result = pytesseract.image_to_string(img)
    logger.info("Image OCR complete: %d chars.", len(result))
    return result
